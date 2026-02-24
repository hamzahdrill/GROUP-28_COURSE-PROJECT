"""
Generate technical report for Group 28's AD-HTC Biorefinery Dashboard.
Output: Group28_AD-HTC_Report.docx
"""

import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def style_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_table_from_data(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()

# Title page
doc.add_paragraph()
title = doc.add_heading('AD-HTC Integrated Biorefinery Dashboard', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Technical Report')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run(
    'MEG 315 — Applied Thermodynamics\n'
    'Group 28\n'
    'Department of Mechanical Engineering\n'
    'University of Lagos\n'
    'February 2026'
)
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# Table of Contents
style_heading(doc, 'Table of Contents')
toc_items = [
    '1. Introduction',
    '2. Theory of the AD-HTC Combined Cycle',
    '3. System Architecture',
    '4. Technology Stack',
    '5. API Endpoints',
    '6. Database Schema',
    '7. Feedstock Database',
    '8. Otto Cycle Analysis',
    '9. Rankine Cycle Analysis',
    '10. Heat Integration',
    '11. Mass & Energy Balance',
    '12. How to Run',
    '13. Conclusion',
    '14. References',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ── 1. INTRODUCTION ──
style_heading(doc, '1. Introduction')
add_para(doc,
    'This report documents the design, implementation, and analysis of an AD-HTC '
    'integrated biorefinery dashboard developed as part of MEG 315 — Applied '
    'Thermodynamics at the University of Lagos. The dashboard simulates the '
    'thermodynamic behaviour of a combined Anaerobic Digestion (AD) and Hydrothermal '
    'Carbonization (HTC) system, including Otto cycle gas power generation, '
    'Rankine steam cycle waste heat recovery, and heat integration analysis.')

add_para(doc,
    'The project features a full-stack web application with a FastAPI backend, '
    'SQLite database for persistent storage, and an interactive HTML/JavaScript '
    'frontend with real-time thermodynamic calculations and visualisations.')

# ── 2. THEORY ──
style_heading(doc, '2. Theory of the AD-HTC Combined Cycle')

style_heading(doc, '2.1 Anaerobic Digestion (AD)', level=2)
add_para(doc,
    'Anaerobic digestion is a biological process in which microorganisms break down '
    'organic matter in the absence of oxygen. Biomass feedstock (such as cow manure, '
    'food waste, or agricultural residues) is fed into a sealed reactor maintained at '
    'mesophilic conditions (35–40°C). The bacteria produce biogas, a mixture of methane '
    '(CH₄, 50–65%) and carbon dioxide (CO₂, 35–50%). The remaining material, called '
    'digestate, is a nutrient-rich slurry.')

style_heading(doc, '2.2 Hydrothermal Carbonization (HTC)', level=2)
add_para(doc,
    'HTC is a thermochemical process that converts wet biomass into hydrochar — a '
    'carbon-rich solid fuel with properties similar to lignite coal. The process '
    'operates at 180–260°C under autogenous pressure (10–40 bar). HTC is particularly '
    'advantageous because it processes wet feedstocks directly, avoiding the energy '
    'cost of drying.')

style_heading(doc, '2.3 Gas Power Cycle (Otto)', level=2)
add_para(doc,
    'The biogas produced by AD is used as fuel in an air-standard Otto cycle '
    '(internal combustion engine). The cycle consists of four processes:\n'
    '• Process 1→2: Isentropic compression — T₂ = T₁ × r^(γ−1)\n'
    '• Process 2→3: Constant-volume heat addition — Q_in = Cp × (T₃ − T₂)\n'
    '• Process 3→4: Isentropic expansion — T₄ = T₃ / r^(γ−1)\n'
    '• Process 4→1: Constant-volume heat rejection\n'
    'The thermal efficiency is η = W_net / Q_in where W_net = W_turb − W_comp.')

style_heading(doc, '2.4 Steam Cycle (Rankine)', level=2)
add_para(doc,
    'The waste heat from the gas turbine exhaust is recovered to generate steam '
    'via a Rankine-type heating loop:\n'
    '• A→B: Pump raises liquid pressure from P_cond to P_boiler (60 bar)\n'
    '• B→C: Boiler heats water to superheated steam using waste heat\n'
    '• C→D: Steam transfers heat to HTC reactor (digestate carbonization)\n'
    '• D→A: Steam condenses and returns to pump\n'
    'Steam properties are obtained from a hardcoded P=6.0 MPa table with '
    'linear interpolation between data points.')

style_heading(doc, '2.5 System Integration', level=2)
add_para(doc,
    'The key innovation is thermal self-sufficiency. Biogas fuels the gas engine, '
    'exhaust waste heat generates steam for HTC, and digestate from AD becomes '
    'the HTC feedstock. A heat integration analysis compares CHP waste heat against '
    'the combined thermal demands of the AD reactor (heating to 37°C) and the HTC '
    'reactor (heating to 180–260°C).')

# ── 3. SYSTEM ARCHITECTURE ──
style_heading(doc, '3. System Architecture')
add_para(doc,
    'The application follows a client-server architecture:')
add_para(doc,
    '• Frontend: Single HTML file (ad-htc-biomass-v2.html) with JavaScript, '
    'Chart.js, and Tailwind CSS\n'
    '• Backend: FastAPI (Python) serving the dashboard and 3 API endpoints\n'
    '• Database: SQLite with two tables (feedstocks, calculations)\n'
    '• The HTML file works both standalone (local calculations) and with the '
    'FastAPI backend (API calls + database storage)')

add_para(doc, '[Insert system architecture diagram here]', italic=True)

# ── 4. TECHNOLOGY STACK ──
style_heading(doc, '4. Technology Stack')
add_table_from_data(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Backend', 'Python / FastAPI', 'RESTful API for cycle calculations'],
        ['Database', 'SQLite', 'Feedstock storage + calculation history'],
        ['Validation', 'Pydantic', 'Input/output data models'],
        ['Frontend', 'HTML / CSS / JS', 'Interactive dashboard UI'],
        ['Charts', 'Chart.js', 'T-s, h-s, energy, mass, heat charts'],
        ['Styling', 'Tailwind CSS', 'Responsive layout'],
        ['Schematic', 'SVG', 'Animated process flow diagram'],
    ]
)

# ── 5. API ENDPOINTS ──
style_heading(doc, '5. API Endpoints')
add_para(doc, 'The FastAPI backend exposes the following RESTful endpoints:')
add_table_from_data(doc,
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/', 'Serve the dashboard HTML page'],
        ['GET', '/api/feedstocks', 'Return all 12 feedstocks from SQLite'],
        ['POST', '/api/analyze', 'Run cycle analysis, store result in DB'],
        ['GET', '/api/history', 'Return past calculation results'],
    ]
)
add_para(doc,
    'The POST /api/analyze endpoint accepts a JSON body with all sidebar '
    'parameters (feedstock_key, feed_rate, T1, compression_ratio, T3, '
    'mass_flow_air, t_steam_C, p_cond, eta_turbine, eta_pump, htc_temp) '
    'and returns a structured JSON response containing otto{}, rankine{}, '
    'mass_balance{}, heat{}, and olr values.')

# ── 6. DATABASE SCHEMA ──
style_heading(doc, '6. Database Schema')
add_para(doc, 'The SQLite database (adhtc.db) contains two tables:')

style_heading(doc, '6.1 feedstocks', level=2)
add_table_from_data(doc,
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['key', 'TEXT UNIQUE', 'Identifier (e.g. cow_manure)'],
        ['name', 'TEXT', 'Display name'],
        ['ts', 'REAL', 'Total solids fraction'],
        ['vs', 'REAL', 'Volatile solids fraction'],
        ['biogas_yield', 'REAL', 'Biogas yield (m³/kg VS)'],
        ['ch4', 'REAL', 'Methane fraction'],
        ['htc_yield', 'REAL', 'Hydrochar mass yield'],
        ['htc_hhv', 'REAL', 'Hydrochar HHV (kJ/kg)'],
    ]
)

style_heading(doc, '6.2 calculations', level=2)
add_table_from_data(doc,
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['timestamp', 'TEXT', 'ISO 8601 timestamp'],
        ['feedstock', 'TEXT', 'Feedstock key used'],
        ['inputs', 'JSON', 'All input parameters'],
        ['otto_results', 'JSON', 'Otto cycle outputs'],
        ['rankine_results', 'JSON', 'Rankine cycle outputs'],
        ['heat_balance', 'REAL', 'Heat surplus/deficit (kW)'],
        ['biogas_vol', 'REAL', 'Biogas volume (m³/day)'],
        ['elec_power', 'REAL', 'Electrical power (kW)'],
        ['hydrochar', 'REAL', 'Hydrochar output (kg/day)'],
    ]
)

# ── 7. FEEDSTOCK DATABASE ──
style_heading(doc, '7. Feedstock Database')
add_para(doc,
    'The database is seeded with 12 real biomass feedstocks with published data:')
add_table_from_data(doc,
    ['Feedstock', 'TS%', 'VS%', 'Biogas (m³/kg VS)', 'CH₄%', 'HTC Yield', 'HTC HHV'],
    [
        ['Cow Manure', '18', '80', '0.280', '60', '48%', '14,500'],
        ['Food / Kitchen Waste', '25', '90', '0.550', '62', '45%', '19,800'],
        ['Sewage Sludge', '30', '70', '0.300', '58', '50%', '15,000'],
        ['Corn Stover', '88', '85', '0.338', '55', '55%', '21,500'],
        ['Rice Husk', '90', '80', '0.250', '52', '58%', '18,200'],
        ['Sugarcane Bagasse', '50', '90', '0.310', '54', '52%', '20,100'],
        ['Pig Manure', '20', '82', '0.320', '62', '46%', '15,200'],
        ['Wood Chips', '85', '92', '0.200', '55', '65%', '24,000'],
        ['Municipal Solid Waste', '60', '75', '0.400', '58', '50%', '17,500'],
        ['Microalgae', '10', '85', '0.450', '65', '40%', '22,000'],
        ['Cassava Peel', '88', '86', '0.420', '58', '48%', '19,200'],
        ['Palm EFB', '78', '85', '0.270', '53', '60%', '20,800'],
    ]
)

# ── 8. OTTO CYCLE ANALYSIS ──
style_heading(doc, '8. Otto Cycle Analysis')

cp_air = 1.005; cp_gas = 1.148; g = 1.4
T1 = 298; r = 10; T3 = 1200
T2 = T1 * r ** (g - 1); T4 = T3 / r ** (g - 1)
Wc = cp_air * (T2 - T1); Wt = cp_gas * (T3 - T4)
Wn = Wt - Wc; Qin = cp_gas * (T3 - T2)
eta = Wn / Qin; elec = 50 * Wn

add_para(doc,
    'Sample calculation with default parameters:\n'
    f'• Intake temperature T₁ = {T1} K\n'
    f'• Compression ratio r = {r}\n'
    f'• Peak temperature T₃ = {T3} K\n'
    f'• Air mass flow ṁ = 50 kg/s')

style_heading(doc, '8.1 State Points', level=2)
add_table_from_data(doc,
    ['Point', 'Process', 'T (K)', 's (kJ/kg·K)'],
    [
        ['1', 'Intake', f'{T1:.1f}', '1.000'],
        ['2', 'Compression', f'{T2:.1f}', '1.000'],
        ['3', 'Combustion', f'{T3:.1f}', f'{1.0+cp_gas*math.log(T3/T2):.3f}'],
        ['4', 'Exhaust', f'{T4:.1f}', f'{1.0+cp_gas*math.log(T3/T2):.3f}'],
    ]
)

style_heading(doc, '8.2 Performance', level=2)
add_para(doc,
    f'• Compression work: W_comp = {Wc:.1f} kJ/kg\n'
    f'• Expansion work: W_turb = {Wt:.1f} kJ/kg\n'
    f'• Net work: W_net = {Wn:.1f} kJ/kg\n'
    f'• Heat input: Q_in = {Qin:.1f} kJ/kg\n'
    f'• Thermal efficiency: η = {eta*100:.1f}%\n'
    f'• Electrical power: {elec:.0f} kW')

add_para(doc, '[Insert T-s diagram screenshot here]', italic=True)

# ── 9. RANKINE CYCLE ──
style_heading(doc, '9. Rankine Cycle Analysis')

hA = 359.9; sA = 1.145; sfg = 7.533 - 1.145; hfg = 2653.6 - 359.9
wp = 0.001 * (60 - 0.06) * 100; hB = hA + wp / 0.80
hC = 3423.1; sC = 6.8826
xs = min(1.0, (sC - 1.145) / sfg)
hDs = hA + xs * hfg; hD = hC - 0.85 * (hC - hDs)
wn_r = (hC - hD) - wp / 0.80; eta_r = wn_r / (hC - hB)

add_para(doc,
    'Steam cycle at P = 6.0 MPa (60 bar) with hardcoded steam table data.\n'
    f'• Steam temperature: 500°C\n'
    f'• Condenser pressure: 0.06 bar\n'
    f'• Turbine efficiency: 0.85\n'
    f'• Pump efficiency: 0.80')

style_heading(doc, '9.1 State Points', level=2)
add_table_from_data(doc,
    ['Point', 'Description', 'h (kJ/kg)', 's (kJ/kg·K)'],
    [
        ['A', 'Pump Inlet', f'{hA:.1f}', f'{sA:.4f}'],
        ['B', 'Pump Outlet', f'{hB:.1f}', f'{sA:.4f}'],
        ['C', 'Turbine Inlet', f'{hC:.1f}', f'{sC:.4f}'],
        ['D', 'Turbine Exit', f'{hD:.1f}', f'{1.145+((hD-hA)/hfg)*sfg:.4f}'],
    ]
)

style_heading(doc, '9.2 Performance', level=2)
add_para(doc,
    f'• Pump work: w_pump = {wp:.2f} kJ/kg\n'
    f'• Net work: w_net = {wn_r:.1f} kJ/kg\n'
    f'• Boiler heat: q_boiler = {hC - hB:.1f} kJ/kg\n'
    f'• Thermal efficiency: η = {eta_r*100:.1f}%')

add_para(doc, '[Insert h-s Mollier diagram screenshot here]', italic=True)

# ── 10. HEAT INTEGRATION ──
style_heading(doc, '10. Heat Integration')
add_para(doc,
    'The heat integration analysis compares available CHP waste heat against '
    'the thermal demands of the AD and HTC reactors:\n\n'
    '• CHP waste heat = ṁ_air × Cp × (T₄ − T₁) × 0.45\n'
    '• AD demand = (feed kg/s) × 4180 × (37°C − ambient) / 1000\n'
    '• HTC demand = (digestate kg/s) × 4000 × (T_HTC − 37°C) / 1000\n'
    '• Heat balance = CHP output − AD demand − HTC demand\n\n'
    'If the balance is positive, the system is self-sufficient (SURPLUS). '
    'If negative, external heating is required (DEFICIT).\n\n'
    'API test result: Heat balance = +4555.89 kW (SURPLUS) for Cow Manure at 10 t/day.')

# ── 11. MASS BALANCE ──
style_heading(doc, '11. Mass & Energy Balance')
feedKg = 10000; TS = 0.18; VS = 0.80
vs = feedKg * TS * VS; bgv = vs * 0.280; ch4 = bgv * 0.60
bgm = bgv * 1.2; dd = vs * 0.40; hc = dd * 0.48

add_para(doc, 'Sample: Cow Manure at 10 tonnes/day')
add_table_from_data(doc,
    ['Output', 'Value', 'Unit'],
    [
        ['Feed input', f'{feedKg}', 'kg/day'],
        ['Volatile solids', f'{vs:.0f}', 'kg/day'],
        ['Biogas volume', f'{bgv:.0f}', 'm³/day'],
        ['Methane volume', f'{ch4:.0f}', 'm³/day'],
        ['Biogas mass', f'{bgm:.0f}', 'kg/day'],
        ['Dry digestate', f'{dd:.0f}', 'kg/day'],
        ['Hydrochar', f'{hc:.0f}', 'kg/day'],
    ]
)

# ── 12. HOW TO RUN ──
style_heading(doc, '12. How to Run')
add_para(doc,
    'Option A — Full-stack (with API + database):\n'
    '  1. pip install -r requirements.txt\n'
    '  2. python main.py\n'
    '  3. Open http://localhost:8000\n'
    '  4. API docs at http://localhost:8000/docs\n\n'
    'Option B — Standalone (no installation):\n'
    '  1. Open ad-htc-biomass-v2.html in any web browser\n'
    '  2. All calculations run locally in JavaScript')

# ── 13. CONCLUSION ──
style_heading(doc, '13. Conclusion')
add_para(doc,
    'This project demonstrates a complete AD-HTC biorefinery simulation with '
    'full-stack web development. Key achievements:\n\n'
    '• FastAPI backend with 4 RESTful API endpoints\n'
    '• SQLite database storing 12 feedstocks and calculation history\n'
    '• Otto cycle analysis with η = 72.6%\n'
    '• Rankine steam cycle with η = 27.7%\n'
    '• Heat integration showing system self-sufficiency\n'
    '• Interactive dashboard with Chart.js visualisations\n'
    '• Zero-installation standalone mode also available\n\n'
    'The project covers key thermodynamic concepts including the First Law, '
    'Second Law, cycle analysis, and heat integration, while demonstrating '
    'modern software engineering practices (RESTful API, SQL, data validation).')

# ── 14. REFERENCES ──
style_heading(doc, '14. References')
refs = [
    'Potrč, S., Petrovič, A., Egieya, J. M., & Čuček, L. (2025). Valorization of biomass through anaerobic digestion and hydrothermal carbonization. Energies, 18(2), 334.',
    'Çengel, Y. A., & Boles, M. A. (2019). Thermodynamics: An Engineering Approach (9th ed.). McGraw-Hill.',
    'FastAPI Documentation. https://fastapi.tiangolo.com/',
    'SQLite Documentation. https://www.sqlite.org/docs.html',
    'Chart.js Documentation. https://www.chartjs.org/docs/',
]
for i, ref in enumerate(refs, 1):
    add_para(doc, f'[{i}] {ref}', size=10)

# ═══════════════════════════════════════════════════════════════
out = os.path.join(SCRIPT_DIR, "Group28_AD-HTC_Report.docx")
doc.save(out)
print(f"Report saved: {out}")
